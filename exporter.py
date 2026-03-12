"""
小说导出模块 - 支持多种格式的小说导出

版权所有 © 2026 新疆幻城网安科技有限责任公司 (幻城科技)
作者：幻城

支持格式:
- Word (.docx) - 专业排版
- 纯文本 (.txt) - 通用格式
- Markdown (.md) - 轻量级标记
- HTML (.html) - 网页格式
"""

import os
import re
import tempfile
import logging
from typing import Tuple, Optional, List, Dict
from datetime import datetime
from pathlib import Path
from src.config.paths import get_exports_dir

logger = logging.getLogger(__name__)

# 导出目录
EXPORT_DIR = get_exports_dir()
EXPORT_DIR.mkdir(exist_ok=True)


def _sanitize_filename(filename: str) -> str:
    """
    清理文件名,移除非法字符

    Args:
        filename: 原始文件名

    Returns:
        清理后的文件名
    """
    # 移除或替换非法字符
    filename = re.sub(r'[<>:"/\\|?*]', '', filename)
    # 限制长度
    if len(filename) > 100:
        filename = filename[:100]
    # 去除首尾空格
    filename = filename.strip()
    return filename or "export"


def _extract_chapters_from_markdown(text: str) -> List[Dict[str, str]]:
    """
    从Markdown文本中提取章节

    Args:
        text: Markdown格式的文本

    Returns:
        章节列表 [{"title": "...", "content": "..."}]
    """
    chapters = []
    lines = text.split('\n')

    current_title = "前言"
    current_content = []

    for line in lines:
        # 检测章节标题 (## 开头)
        if line.startswith('## '):
            # 保存上一章
            if current_content:
                chapters.append({
                    "title": current_title,
                    "content": '\n'.join(current_content).strip()
                })

            # 开始新章节
            current_title = line[3:].strip()
            current_content = []
        else:
            current_content.append(line)

    # 保存最后一章
    if current_content:
        chapters.append({
            "title": current_title,
            "content": '\n'.join(current_content).strip()
        })

    return chapters


def export_to_txt(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为纯文本格式

    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题

    Returns:
        (文件路径, 状态信息)
    """
    logger.info(f"[导出功能] 开始导出TXT: {title}, 内容长度: {len(novel_text)}")

    try:
        if not novel_text.strip():
            return None, "无内容可导出"

        # 清理标题
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        filepath = EXPORT_DIR / filename

        # 写入文件（原子写入）
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.txt', dir=EXPORT_DIR)
            os.close(tmp_fd)

            with open(tmp_path, 'w', encoding='utf-8') as f:
                f.write(f"{title}\n")
                f.write(f"生成日期: {datetime.now().strftime('%Y年%m月%d日')}\n")
                f.write("=" * 50 + "\n\n")
                f.write(novel_text)

            os.replace(tmp_path, str(filepath))

        except Exception as e:
            logger.error(f"写入TXT文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"[导出功能] TXT导出成功: {filename}")
        return str(filepath), f"导出成功: {filename}"

    except Exception as e:
        logger.error(f"[导出功能] TXT导出失败: {e}", exc_info=True)
        return None, f"导出失败: {str(e)}"


def export_to_markdown(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为Markdown格式

    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题

    Returns:
        (文件路径, 状态信息)
    """
    logger.info(f"[导出功能] 开始导出Markdown: {title}, 内容长度: {len(novel_text)}")

    try:
        if not novel_text.strip():
            return None, "无内容可导出"

        # 清理标题
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.md"
        filepath = EXPORT_DIR / filename

        # 写入文件（原子写入）
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.md', dir=EXPORT_DIR)
            os.close(tmp_fd)

            with open(tmp_path, 'w', encoding='utf-8') as f:
                f.write(f"# {title}\n\n")
                f.write(f"**生成日期**: {datetime.now().strftime('%Y年%m月%d日')}\n\n")
                f.write("---\n\n")
                f.write(novel_text)

            os.replace(tmp_path, str(filepath))

        except Exception as e:
            logger.error(f"写入MD文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"[导出功能] Markdown导出成功: {filename}")
        return str(filepath), f"导出成功: {filename}"

    except Exception as e:
        logger.error(f"[导出功能] Markdown导出失败: {e}", exc_info=True)
        return None, f"导出失败: {str(e)}"


def export_to_html(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为HTML格式

    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题

    Returns:
        (文件路径, 状态信息)
    """
    logger.info(f"[导出功能] 开始导出HTML: {title}, 内容长度: {len(novel_text)}")

    try:
        if not novel_text.strip():
            return None, "无内容可导出"

        # 清理标题
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.html"
        filepath = EXPORT_DIR / filename

        # 简单的Markdown转HTML
        chapters = _extract_chapters_from_markdown(novel_text)

        html_content = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>{title}</title>
    <style>
        body {{
            max-width: 800px;
            margin: 0 auto;
            padding: 20px;
            font-family: "Microsoft YaHei", "SimHei", sans-serif;
            line-height: 1.8;
            color: #333;
        }}
        h1 {{
            text-align: center;
            color: #2c3e50;
            border-bottom: 2px solid #3498db;
            padding-bottom: 10px;
        }}
        h2 {{
            color: #34495e;
            margin-top: 30px;
            padding-bottom: 5px;
            border-bottom: 1px solid #bdc3c7;
        }}
        .meta {{
            text-align: center;
            color: #7f8c8d;
            margin-bottom: 30px;
        }}
        p {{
            text-indent: 2em;
            margin: 10px 0;
        }}
    </style>
</head>
<body>
    <h1>{title}</h1>
    <div class="meta">生成日期: {datetime.now().strftime('%Y年%m月%d日')}</div>
"""

        # 添加章节内容
        for chapter in chapters:
            html_content += f"    <h2>{chapter['title']}</h2>\n"
            paragraphs = chapter['content'].split('\n\n')
            for para in paragraphs:
                if para.strip():
                    # 转义HTML特殊字符
                    escaped = para.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')
                    html_content += f"    <p>{escaped}</p>\n"

        html_content += """
</body>
</html>
"""

        # 写入文件（原子写入）
        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.html', dir=EXPORT_DIR)
            os.close(tmp_fd)

            with open(tmp_path, 'w', encoding='utf-8') as f:
                f.write(html_content)

            os.replace(tmp_path, str(filepath))

        except Exception as e:
            logger.error(f"写入HTML文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"[导出功能] HTML导出成功: {filename}")
        return str(filepath), f"导出成功: {filename}"

    except Exception as e:
        logger.error(f"[导出功能] HTML导出失败: {e}", exc_info=True)
        return None, f"导出失败: {str(e)}"


def export_to_docx(novel_text: str, title: str) -> Tuple[Optional[str], str]:
    """
    导出为Word (DOCX) 格式 - 专业排版

    Args:
        novel_text: 小说文本（Markdown格式）
        title: 小说标题

    Returns:
        (文件路径, 状态信息)
    """
    logger.info(f"[导出功能] 开始导出DOCX: {title}, 内容长度: {len(novel_text)}")

    try:
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.oxml.ns import qn
    except ImportError:
        return None, "导出Word文档需要安装python-docx: pip install python-docx"

    try:
        if not novel_text.strip():
            return None, "无内容可导出"

        # 提取章节
        chapters = _extract_chapters_from_markdown(novel_text)

        if not chapters:
            return None, "无法从文本中提取章节"

        doc = Document()

        # 配置样式
        style = doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)

        # 中文字体
        rPr = style.element.get_or_add_rPr()
        rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '宋体')

        # 段落格式
        style.paragraph_format.first_line_indent = Pt(24)
        style.paragraph_format.space_after = Pt(0)
        style.paragraph_format.line_spacing = 1.5

        # 添加书名
        title_para = doc.add_paragraph(title)
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title_para.runs[0]
        title_run.font.name = '黑体'
        title_run.font.size = Pt(26)
        title_run.font.bold = True

        # 中文字体设置
        title_rPr = title_run._element.get_or_add_rPr()
        title_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '黑体')

        # 添加作者和日期信息
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(
            f"生成日期：{datetime.now().strftime('%Y年%m月%d日')}"
        )
        info_run.font.size = Pt(10)

        doc.add_paragraph()  # 空行

        # 添加章节
        for chapter in chapters:
            # 章节标题
            chapter_title_para = doc.add_paragraph(chapter['title'])
            chapter_title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

            for run in chapter_title_para.runs:
                run.font.name = '黑体'
                run.font.size = Pt(16)
                run.font.bold = True
                run_rPr = run._element.get_or_add_rPr()
                run_rPr.find(qn('w:rFonts')).set(qn('w:eastAsia'), '黑体')

            doc.add_paragraph()  # 空行

            # 章节内容 - 按段落添加
            paragraphs = chapter['content'].split('\n\n')
            for para_text in paragraphs:
                if para_text.strip():
                    p = doc.add_paragraph(para_text.strip(), style='Normal')

            doc.add_paragraph()  # 章节间空行

        # 保存文件（原子写入）
        safe_title = _sanitize_filename(title)
        filename = f"{safe_title}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        filepath = EXPORT_DIR / filename

        try:
            tmp_fd, tmp_path = tempfile.mkstemp(suffix='.docx', dir=EXPORT_DIR)
            os.close(tmp_fd)
            doc.save(tmp_path)
            os.replace(tmp_path, str(filepath))
        except Exception as e:
            logger.error(f"写入DOCX文件失败: {e}")
            return None, f"导出失败: {e}"

        logger.info(f"[导出功能] DOCX导出成功: {filename}")
        return str(filepath), f"导出成功: {filename}"

    except Exception as e:
        logger.error(f"[导出功能] DOCX导出失败: {e}", exc_info=True)
        return None, f"导出失败: {str(e)}"
