"""
项目管理模块 - 支持保存、加载、导出项目，集成 Supabase 云端同步

版权所有 © 2026 新疆幻城网安科技有限责任公司 (幻城科技)
作者：幻城
"""
import json
import os
import re
import tempfile
import logging
import threading
from typing import List, Dict, Optional, Tuple
from datetime import datetime
from pathlib import Path
from src.config.paths import get_projects_dir
from src.config.supabase_client import (
    ProjectSyncManager,
    is_sync_enabled,
    restore_projects_from_cloud
)

logger = logging.getLogger(__name__)

PROJECTS_DIR = get_projects_dir()
PROJECTS_DIR.mkdir(parents=True, exist_ok=True)

# 全局同步管理器
_project_sync_manager: Optional[ProjectSyncManager] = None
_sync_lock = threading.Lock()


def get_project_sync_manager() -> Optional[ProjectSyncManager]:
    """获取项目同步管理器实例"""
    global _project_sync_manager
    if _project_sync_manager is None and is_sync_enabled():
        _project_sync_manager = ProjectSyncManager()
    return _project_sync_manager


class ProjectManager:
    """项目管理器"""

    @staticmethod
    def _slugify(name: str) -> str:
        s = str(name or "").lower()
        s = re.sub(r'[^a-z0-9]+', '-', s)
        s = s.strip('-')
        return s or 'project'

    @staticmethod
    def create_project(
        title: str,
        genre: str,
        character_setting: str,
        world_setting: str,
        plot_idea: str,
        chapter_count: int = 50
    ) -> Tuple[Optional[str], str]:
        """
        创建新项目

        Returns:
            (项目ID, 状态信息)
        """
        if not title or not title.strip():
            return None, "项目标题不能为空"

        if not genre or not genre.strip():
            return None, "小说类型不能为空"

        # 生成项目ID
        slug = ProjectManager._slugify(title.strip())
        project_id = f"{slug}_{datetime.now().strftime('%Y%m%d_%H%M%S')}"

        # 创建项目数据
        project_data = {
            "id": project_id,
            "title": title.strip(),
            "genre": genre.strip(),
            "character_setting": character_setting.strip(),
            "world_setting": world_setting.strip(),
            "plot_idea": plot_idea.strip(),
            "chapter_count": chapter_count,
            "chapters": [],
            "outline": [],
            "created_at": datetime.now().isoformat(),
            "updated_at": datetime.now().isoformat()
        }

        # 保存项目
        try:
            project_file = Path(PROJECTS_DIR) / f"{project_id}.json"
            with open(project_file, 'w', encoding='utf-8') as f:
                json.dump(project_data, f, ensure_ascii=False, indent=2)

            # 同步到云端（异步执行，不阻塞）
            def sync_to_cloud():
                try:
                    sync_mgr = get_project_sync_manager()
                    if sync_mgr:
                        sync_mgr.sync_project(project_id, project_data)
                except Exception as e:
                    logger.warning(f"[项目管理] 同步项目到云端失败: {e}")

            # 在后台线程中同步
            threading.Thread(target=sync_to_cloud, daemon=True).start()

            logger.info(f"[项目管理] 创建项目: {title} (ID: {project_id}), 类型: {genre}, 章节数: {chapter_count}")
            return project_id, f"项目创建成功: {project_id}"

        except Exception as e:
            logger.error(f"[项目管理] 项目创建失败: {e}", exc_info=True)
            return None, f"项目创建失败: {str(e)}"

    @staticmethod
    def save_project(project_id: str, project_data: Dict) -> Tuple[bool, str]:
        """
        保存项目到磁盘，并同步到云端

        Args:
            project_id: 项目ID
            project_data: 项目数据字典

        Returns:
            (成功标志, 状态信息)
        """
        try:
            if not project_data:
                return False, "项目数据为空"

            # 更新时间戳
            project_data["updated_at"] = datetime.now().isoformat()

            # 保存到文件
            project_file = Path(PROJECTS_DIR) / f"{project_id}.json"
            with open(project_file, 'w', encoding='utf-8') as f:
                json.dump(project_data, f, ensure_ascii=False, indent=2)

            chapter_count = len(project_data.get("chapters", []))
            logger.info(f"[项目管理] 项目已保存: {project_id}, 章节数: {chapter_count}")

            # 同步到云端（异步执行，不阻塞）
            def sync_to_cloud():
                try:
                    sync_mgr = get_project_sync_manager()
                    if sync_mgr:
                        sync_mgr.sync_project(project_id, project_data)
                        logger.debug(f"[项目管理] 项目已同步到云端: {project_id}")
                except Exception as e:
                    logger.warning(f"[项目管理] 同步项目到云端失败: {e}")

            # 在后台线程中同步
            threading.Thread(target=sync_to_cloud, daemon=True).start()

            return True, f"项目已保存: {project_id}"

        except Exception as e:
            logger.error(f"[项目管理] 项目保存失败: {e}", exc_info=True)
            return False, f"项目保存失败: {str(e)}"

    @staticmethod
    def get_project(project_id: str) -> Optional[Dict]:
        """
        获取项目数据（支持新旧格式）

        Returns:
            项目数据字典 或 None
        """
        try:
            # 尝试新格式（单文件）
            project_file = Path(PROJECTS_DIR) / f"{project_id}.json"

            if project_file.exists():
                with open(project_file, 'r', encoding='utf-8') as f:
                    project_data = json.load(f)
                logger.info(f"[项目管理] 项目已加载（新格式）: {project_id}")
                return project_data

            # 尝试旧格式（目录）
            project_dir = Path(PROJECTS_DIR) / project_id
            if project_dir.exists():
                metadata_file = project_dir / "metadata.json"
                if metadata_file.exists():
                    with open(metadata_file, 'r', encoding='utf-8') as f:
                        old_data = json.load(f)

                    # 迁移到新格式
                    logger.info(f"[项目管理] 检测到旧格式项目，正在迁移: {project_id}")
                    new_project_data = {
                        "id": project_id,
                        "title": old_data.get("title", ""),
                        "genre": old_data.get("genre", ""),
                        "character_setting": old_data.get("character_setting", ""),
                        "world_setting": old_data.get("world_setting", ""),
                        "plot_idea": old_data.get("plot_idea", ""),
                        "chapter_count": old_data.get("total_chapters", len(old_data.get("chapters", []))),
                        "chapters": old_data.get("chapters", []),
                        "outline": old_data.get("outline", []),
                        "created_at": old_data.get("created_at", ""),
                        "updated_at": old_data.get("updated_at", "")
                    }

                    # 保存为新格式
                    ProjectManager.save_project(project_id, new_project_data)

                    logger.info(f"[项目管理] 项目已迁移并加载: {project_id}")
                    return new_project_data

            logger.warning(f"[项目管理] 项目不存在: {project_id}")
            return None

        except Exception as e:
            logger.error(f"[项目管理] 项目加载失败: {e}", exc_info=True)
            return None

    @staticmethod
    def list_projects() -> List[Dict]:
        """
        列出所有项目

        Returns:
            项目信息列表 [{"id": "...", "title": "...", "genre": "...", "created_at": "...", "updated_at": "...", "chapter_count": ...}]
        """
        try:
            projects = []
            projects_path = Path(PROJECTS_DIR)

            if not projects_path.exists():
                return projects

            # 遍历所有JSON文件
            for project_file in projects_path.glob("*.json"):
                try:
                    with open(project_file, 'r', encoding='utf-8') as f:
                        project_data = json.load(f)

                    chapters = project_data.get("chapters", [])
                    completed_count = sum(1 for ch in chapters if ch.get("content", "").strip())

                    projects.append({
                        "id": project_file.stem,
                        "title": project_data.get("title", "未命名"),
                        "genre": project_data.get("genre", ""),
                        "created_at": project_data.get("created_at", ""),
                        "updated_at": project_data.get("updated_at", ""),
                        "total_chapters": project_data.get("chapter_count", len(chapters)),
                        "chapter_count": len(chapters),
                        "completed_chapters": completed_count
                    })
                except Exception as e:
                    logger.warning(f"[项目管理] 读取项目文件失败 {project_file.name}: {e}")

            # 按更新时间排序
            projects.sort(key=lambda x: x.get("updated_at", ""), reverse=True)

            logger.info(f"[项目管理] 找到 {len(projects)} 个项目")
            return projects

        except Exception as e:
            logger.error(f"[项目管理] 列出项目失败: {e}", exc_info=True)
            return []

    @staticmethod
    def get_project_by_title(project_title: str) -> Optional[Dict]:
        """
        按标题获取项目信息

        Returns:
            项目字典 或 None
        """
        projects = ProjectManager.list_projects()
        for project in projects:
            if project.get("title") == project_title:
                # 返回完整的项目数据
                return ProjectManager.get_project(project["id"])
        return None

    @staticmethod
    def delete_project(project_id: str) -> Tuple[bool, str]:
        """
        删除项目（本地和云端）

        Returns:
            (成功标志, 状态信息)
        """
        try:
            project_file = Path(PROJECTS_DIR) / f"{project_id}.json"

            if not project_file.exists():
                # 仍然尝试删除云端数据
                sync_mgr = get_project_sync_manager()
                if sync_mgr:
                    sync_mgr.delete_project(project_id)
                return False, f"项目不存在: {project_id}"

            project_file.unlink()

            # 同时删除云端数据（异步）
            def delete_from_cloud():
                try:
                    sync_mgr = get_project_sync_manager()
                    if sync_mgr:
                        sync_mgr.delete_project(project_id)
                        logger.debug(f"[项目管理] 项目已从云端删除: {project_id}")
                except Exception as e:
                    logger.warning(f"[项目管理] 从云端删除项目失败: {e}")

            threading.Thread(target=delete_from_cloud, daemon=True).start()

            logger.info(f"[项目管理] 项目已删除: {project_id}")
            return True, f"项目已删除: {project_id}"

        except Exception as e:
            logger.error(f"[项目管理] 项目删除失败: {e}", exc_info=True)
            return False, f"项目删除失败: {str(e)}"

    @staticmethod
    def export_project(project_id: str, export_format: str = "json") -> Tuple[Optional[str], str]:
        """
        导出项目

        Args:
            project_id: 项目ID
            export_format: 导出格式 (json/txt/md/docx)

        Returns:
            (文件路径, 状态信息)
        """
        try:
            project_data = ProjectManager.get_project(project_id)
            if not project_data:
                return None, f"项目不存在: {project_id}"

            export_dir = Path(PROJECTS_DIR) / "exports"
            export_dir.mkdir(exist_ok=True)

            if export_format == "json":
                filename = f"{project_data.get('title', project_id)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json"
                filepath = export_dir / filename

                with open(filepath, 'w', encoding='utf-8') as f:
                    json.dump(project_data, f, ensure_ascii=False, indent=2)

                logger.info(f"[项目管理] 项目已导出: {filename}, 格式: {export_format}")
                return str(filepath), f"项目已导出: {filename}"

            elif export_format in ["txt", "md"]:
                # 导出为文本
                content = f"# {project_data.get('title', '')}\n\n"
                content += f"类型: {project_data.get('genre', '')}\n"
                content += f"创建时间: {project_data.get('created_at', '')}\n\n"
                content += "## 章节内容\n\n"

                chapters = project_data.get("chapters", [])
                for ch in chapters:
                    if ch.get("content"):
                        content += f"### 第{ch.get('num', '')}章 {ch.get('title', '')}\n\n"
                        content += ch.get("content", "") + "\n\n"

                ext = "md" if export_format == "md" else "txt"
                filename = f"{project_data.get('title', project_id)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.{ext}"
                filepath = export_dir / filename

                with open(filepath, 'w', encoding='utf-8') as f:
                    f.write(content)

                logger.info(f"[项目管理] 项目已导出: {filename}, 格式: {export_format}")
                return str(filepath), f"项目已导出: {filename}"

            elif export_format == "docx":
                try:
                    from docx import Document

                    doc = Document()
                    doc.add_heading(project_data.get('title', ''), 0)
                    doc.add_paragraph(f"类型: {project_data.get('genre', '')}")

                    chapters = project_data.get("chapters", [])
                    for ch in chapters:
                        if ch.get("content"):
                            doc.add_heading(f"第{ch.get('num', '')}章 {ch.get('title', '')}", 1)
                            doc.add_paragraph(ch.get("content", ""))

                    filename = f"{project_data.get('title', project_id)}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
                    filepath = export_dir / filename
                    doc.save(filepath)

                    logger.info(f"[项目管理] 项目已导出: {filename}, 格式: docx")
                    return str(filepath), f"项目已导出: {filename}"

                except ImportError:
                    return None, "导出Word文档需要安装 python-docx: pip install python-docx"
                except Exception as e:
                    return None, f"导出Word文档失败: {str(e)}"
            else:
                logger.warning(f"[项目管理] 不支持的导出格式: {export_format}")
                return None, f"不支持的导出格式: {export_format}"

        except Exception as e:
            logger.error(f"[项目管理] 项目导出失败: {e}", exc_info=True)
            return None, f"项目导出失败: {str(e)}"


def get_project_manager() -> ProjectManager:
    """获取项目管理器实例"""
    return ProjectManager()
